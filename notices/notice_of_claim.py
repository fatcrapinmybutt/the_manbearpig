import os
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'FEDERAL', 'HOUSING')


def build_notice():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading('Notice of Claim – 42 USC §1983', 0)
    body = [
        'Recipients: Muskegon County, Shady Oaks Park MHP LLC, Homes of America LLC, CT Corporation',
        'Basis: unlawful entry and destruction of titled home during writ execution',
        'Utility shutoff and fraudulent billing',
        'Violation of due process under the Fourteenth Amendment'
    ]
    for line in body:
        doc.add_paragraph(line)
    out_path = os.path.join(OUTPUT_DIR, 'notice_of_claim_section1983_shady_oaks.docx')
    doc.save(out_path)
    print(f'Notice of claim saved to {out_path}')


if __name__ == '__main__':
    build_notice()
