import os
import json
import zipfile
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'FOIA')
LOG_PATH = os.path.join(OUTPUT_DIR, 'foia_request_log.json')

REQUESTS = {
    'EGLE': {
        'filename': 'EGLE_request.docx',
        'body': [
            'Subject: Sewer leak reports for 1977 Whitehall Rd (Feb–June 2025).',
            'Please provide all records of inspections, violation notices, and communications with the park owners.'
        ]
    },
    'County_Clerk': {
        'filename': 'County_Clerk_request.docx',
        'body': [
            'Subject: Policies on counterclaims in summary proceedings.',
            'Please provide any rejection notices or memos related to filings labelled "moot" since May 2025.'
        ]
    },
    'Sheriff': {
        'filename': 'Sheriff_request.docx',
        'body': [
            'Subject: Execution of eviction writ for Lot 17, 1977 Whitehall Rd.',
            'Request any logs or bodycam footage documenting the writ execution.'
        ]
    },
}

def build_requests():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    log = []
    for key, info in REQUESTS.items():
        doc = Document()
        doc.add_heading(f'FOIA Request – {key}', 0)
        for para in info['body']:
            doc.add_paragraph(para)
        out_path = os.path.join(OUTPUT_DIR, info['filename'])
        doc.save(out_path)
        log.append({'agency': key, 'file': info['filename']})
    with open(LOG_PATH, 'w') as f:
        json.dump(log, f, indent=2)
    zip_path = os.path.join(OUTPUT_DIR, 'FOIA_PACKET_SHADY_OAKS_2025.zip')
    with zipfile.ZipFile(zip_path, 'w') as z:
        for item in log:
            z.write(os.path.join(OUTPUT_DIR, item['file']), item['file'])
    print(f'FOIA packet created at {zip_path}')

if __name__ == '__main__':
    build_requests()
