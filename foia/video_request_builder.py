import os
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'FOIA')


def build_video_request():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading('FOIA Request – Sheriff Video', 0)
    body = [
        'Department: Muskegon County Sheriff',
        'Time Window: May 23–May 26, 2025',
        'Subject: Writ execution at 1977 Whitehall Rd, Lot 17',
        'Request: bodycam footage, entry logs, statements, dispatch audio'
    ]
    for line in body:
        doc.add_paragraph(line)
    out_path = os.path.join(OUTPUT_DIR, 'foia_request_sheriff_entry_footage.docx')
    doc.save(out_path)
    print(f'Sheriff FOIA request saved to {out_path}')


if __name__ == '__main__':
    build_video_request()
