import os
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'MOTIONS', 'HOUSING')
DOCX_PATH = os.path.join(OUTPUT_DIR, 'emergency_injunction_motion_shady_oaks.docx')


def generate_motion():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading('Emergency Motion for Preliminary Injunction', 0)
    points = [
        'Irreparable harm: home destroyed, utilities shut off',
        'Likelihood of success on the merits: billing contradictions and entity fraud',
        'Public interest favors stopping fraudulent landlord conduct',
        'No adequate remedy at law given destruction of property',
    ]
    for p in points:
        doc.add_paragraph(p, style='List Bullet')
    doc.save(DOCX_PATH)
    print(f'Emergency motion saved to {DOCX_PATH}')


if __name__ == '__main__':
    generate_motion()
