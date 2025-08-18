import os
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'MOTIONS', 'HOUSING')

def build_protective_order():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading('Emergency Motion for Protective Order', 0)
    doc.add_paragraph('Relief requested: prohibit further entry and billing until fraud claims are resolved.')
    doc.add_paragraph('Authority: MCR 2.302(C); Zurcher v. Stanford Daily, 436 U.S. 547.')
    out_path = os.path.join(OUTPUT_DIR, 'motion_protective_order_postwrit_entry.docx')
    doc.save(out_path)
    print(f'Protective order motion saved to {out_path}')

if __name__ == '__main__':
    build_protective_order()
