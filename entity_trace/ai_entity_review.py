import os
import json
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'ENTITY_TRACE')
JSON_PATH = os.path.join(OUTPUT_DIR, 'ai_entity_review_report_HOA_Alden.json')
DOCX_PATH = os.path.join(OUTPUT_DIR, 'entity_overlap_matrix.docx')

ENTITY_DATA = [
    {
        "entity": "Shady Oaks Park MHP LLC",
        "source": "LARA",
        "overlap": "Same RA (C.T. Corp), uses Homes of America billing",
        "flag": True,
    },
    {
        "entity": "Shady Oaks MHP LLC",
        "source": "LARA",
        "overlap": "RA also C.T. Corp, same mailing address",
        "flag": True,
    },
    {
        "entity": "Homes of America LLC",
        "source": "FOIA + LARA",
        "overlap": "Handles Zego billing for park",
        "flag": True,
    },
    {
        "entity": "Alden Global Capital",
        "source": "SEC",
        "overlap": "Investment control of residential REITs including HOA",
        "flag": False,
    },
    {
        "entity": "Cricklewood MHP LLC",
        "source": "LARA",
        "overlap": "Shared manager Kimberly Davis",
        "flag": True,
    },
]


def generate_reports():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    with open(JSON_PATH, 'w') as f:
        json.dump(ENTITY_DATA, f, indent=2)

    doc = Document()
    doc.add_heading('Entity Overlap Matrix', 0)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Entity'
    hdr_cells[1].text = 'Filing Source'
    hdr_cells[2].text = 'Overlap Detected'
    hdr_cells[3].text = 'Flag'
    for row in ENTITY_DATA:
        cells = table.add_row().cells
        cells[0].text = row['entity']
        cells[1].text = row['source']
        cells[2].text = row['overlap']
        cells[3].text = '✅' if row['flag'] else '🔶'
    doc.save(DOCX_PATH)
    print(f'Report written to {JSON_PATH} and {DOCX_PATH}')


if __name__ == '__main__':
    generate_reports()
