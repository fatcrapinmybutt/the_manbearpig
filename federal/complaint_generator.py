import os
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'FEDERAL', 'HOUSING')
DOCX_PATH = os.path.join(OUTPUT_DIR, 'federal_complaint_section1983_conversion.docx')


def generate_complaint():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading('Federal Complaint', 0)
    doc.add_paragraph('Court: U.S. District Court – WDMI')
    doc.add_paragraph('Plaintiff: Andrew Pigors')
    doc.add_paragraph('Defendants: Shady Oaks Park MHP LLC, Homes of America LLC, John Doe Owners')

    doc.add_heading('Counts', level=1)
    counts = [
        'Count I – Deprivation of Property Without Due Process (42 USC §1983)',
        'Count II – Conversion (MCL 750.362)',
        'Count III – Fraudulent Misrepresentation',
        'Count IV – Civil Conspiracy',
        'Count V – Constructive Eviction / Retaliatory Eviction',
    ]
    for c in counts:
        doc.add_paragraph(c, style='List Number')

    doc.add_heading('Key Allegations', level=1)
    allegations = [
        'Rent paid through April; home destroyed post-writ',
        'No lawful possession order over titled home',
        'Water shutoff with child present',
        'Utility billing continued during uninhabitable conditions',
        'Court failed to address entity mismatch',
    ]
    for a in allegations:
        doc.add_paragraph(a, style='List Bullet')

    doc.save(DOCX_PATH)
    print(f'Federal complaint saved to {DOCX_PATH}')


if __name__ == '__main__':
    generate_complaint()
