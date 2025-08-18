import os
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'VIOLATIONS')
DOCX_PATH = os.path.join(OUTPUT_DIR, 'judicial_misconduct_letter_60th_canon_jtc.docx')


def generate_letter():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading('Judicial Misconduct Letter', 0)
    doc.add_paragraph('To: Chief Judge of the 60th District Court and Michigan Judicial Tenure Commission')
    doc.add_paragraph('')
    bullet_points = [
        'Judge proceeded with eviction despite conditional trial agreement not being honored',
        'Clerk refused counterclaim filings under false claim of mootness',
        'Utility shutoffs ignored during eviction',
        'Writ executed for lot only, home entered and destroyed',
    ]
    doc.add_heading('Allegations', level=1)
    for pt in bullet_points:
        doc.add_paragraph(pt, style='List Bullet')

    doc.add_heading('Canons Cited', level=1)
    doc.add_paragraph('Canon 2(A) – Impropriety and appearance of impropriety')
    doc.add_paragraph('Canon 3(B)(8) – Failure to give parties a fair opportunity')
    doc.add_paragraph('Canon 3(C)(1)(a) – Appearance of bias')

    doc.save(DOCX_PATH)
    print(f'Misconduct letter saved to {DOCX_PATH}')


if __name__ == '__main__':
    generate_letter()
