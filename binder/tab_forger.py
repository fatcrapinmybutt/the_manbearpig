import os
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'BINDER_TABS')

EXHIBITS = [
    {'label': 'C-1', 'description': 'Photo – Home with door removed'},
    {'label': 'C-2', 'description': 'May 15 rent and utility bill'},
    {'label': 'C-3', 'description': 'Screenshot child present during shutoff'},
    {'label': 'C-4', 'description': 'Sewer overflow photos'},
    {'label': 'C-5', 'description': 'Clerk email rejecting filing'},
]


def forge_tab():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading('Binder Tab C – Post-Writ Evidence', 0)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Exhibit'
    hdr_cells[1].text = 'Description'
    for ex in EXHIBITS:
        row_cells = table.add_row().cells
        row_cells[0].text = ex['label']
        row_cells[1].text = ex['description']
    out_path = os.path.join(OUTPUT_DIR, 'Binder_Tab_C_PostWrit_Evidence.docx')
    doc.save(out_path)
    print(f'Binder Tab saved to {out_path}')


if __name__ == '__main__':
    forge_tab()
