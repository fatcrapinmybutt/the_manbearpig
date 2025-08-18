import os
import datetime
from docx import Document

FORM_TEMPLATE = "MBP_SYSTEM/FORMS/MC_230_TEMPLATE.docx"
OUTPUT_DIR = "MBP_SYSTEM/VFS/GENERATED_ZIPS"


def build_motion(motion_type, facts, relief):
    if not os.path.exists(FORM_TEMPLATE):
        print("Missing form template.")
        return

    doc = Document(FORM_TEMPLATE)
    doc.add_paragraph(f"Motion Type: {motion_type}")
    doc.add_paragraph(f"Facts: {facts}")
    doc.add_paragraph(f"Requested Relief: {relief}")

    filename = f"Motion_{motion_type.replace(' ', '_')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"Motion generated: {path}")


if __name__ == "__main__":
    build_motion(
        "Parenting Time Enforcement",
        "Exchange denial on 6/1/25",
        "Immediate enforcement and make-up time",
    )
