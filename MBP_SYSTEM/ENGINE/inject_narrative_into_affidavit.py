from docx import Document
import os
import datetime

SOURCE = 'MBP_SYSTEM/VFS/GENERATED_ZIPS/Chronological_Narrative.txt'
TEMPLATE = 'MBP_SYSTEM/FORMS/AFFIDAVIT_TEMPLATE.docx'
OUTFILE = f"MBP_SYSTEM/VFS/GENERATED_ZIPS/Affidavit_Chrono_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"


def inject():
    if not os.path.exists(SOURCE):
        print('Chronological Narrative not found.')
        return

    if not os.path.exists(TEMPLATE):
        print('Affidavit template not found.')
        return

    with open(SOURCE, 'r', encoding='utf-8') as f:
        timeline = f.read()

    doc = Document(TEMPLATE)
    doc.add_heading('Chronological Timeline', level=1)
    doc.add_paragraph(timeline)
    doc.save(OUTFILE)
    print(f'Timeline injected into affidavit: {OUTFILE}')


if __name__ == '__main__':
    inject()
