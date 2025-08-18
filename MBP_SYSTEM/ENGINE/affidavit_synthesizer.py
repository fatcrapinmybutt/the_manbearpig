import os
import datetime
from docx import Document


def synthesize_affidavit(name, timeline, statement, attached_exhibits=None):
    if attached_exhibits is None:
        attached_exhibits = []

    doc = Document()
    doc.add_heading(f"Affidavit of {name}", 0)
    doc.add_paragraph(f"I, {name}, state the following under oath:")
    doc.add_paragraph(statement)

    doc.add_heading("Timeline of Events", level=1)
    for entry in timeline:
        doc.add_paragraph(f"- {entry}")

    if attached_exhibits:
        doc.add_heading("Attached Exhibits", level=1)
        for ex in attached_exhibits:
            doc.add_paragraph(f"Exhibit {ex}")

    filename = f"Affidavit_{name}_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    output_path = os.path.join("MBP_SYSTEM/VFS/GENERATED_ZIPS", filename)
    doc.save(output_path)
    print(f"Affidavit saved to: {output_path}")


if __name__ == "__main__":
    synthesize_affidavit(
        "Andrew J Pigors",
        ["3/26/24: Parenting time denied", "3/29/24: False welfare check filed"],
        "This affidavit is submitted in support of my motion to modify custody.",
        ["A", "B"],
    )
