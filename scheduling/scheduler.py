import os
import json
from datetime import datetime
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'SCHEDULING')
ICS_PATH = os.path.join(OUTPUT_DIR, 'housing_case_timeline.ics')
DOCX_PATH = os.path.join(OUTPUT_DIR, 'court_calendar_printable.docx')
TIMELINE_FILE = os.path.join('data', 'timeline.json')


def load_events():
    if os.path.exists(TIMELINE_FILE):
        with open(TIMELINE_FILE, 'r') as f:
            return json.load(f)
    return []


def export_ics(events):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    lines = ["BEGIN:VCALENDAR", "VERSION:2.0"]
    for e in events:
        dt = datetime.fromisoformat(e['date'])
        lines.append("BEGIN:VEVENT")
        lines.append(f"SUMMARY:{e['description']}")
        lines.append(f"DTSTART:{dt.strftime('%Y%m%dT%H%M%S')}")
        lines.append("END:VEVENT")
    lines.append("END:VCALENDAR")
    with open(ICS_PATH, 'w') as f:
        f.write('\n'.join(lines))
    print(f'ICS calendar saved to {ICS_PATH}')


def export_docx(events):
    doc = Document()
    doc.add_heading('Court Calendar', 0)
    for e in events:
        doc.add_paragraph(f"{e['date']} - {e['description']}")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f'Printable calendar saved to {DOCX_PATH}')


def build_schedule():
    events = load_events()
    export_ics(events)
    export_docx(events)


if __name__ == '__main__':
    build_schedule()
