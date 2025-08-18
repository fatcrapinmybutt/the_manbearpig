import os
from docx import Document

BASE_DIR = os.getenv('LEGAL_RESULTS_DIR', os.path.join('F:/', 'LegalResults'))
OUTPUT_DIR = os.path.join(BASE_DIR, 'PUBLIC_RELEASE')

SUMMARY = (
    "Trailer Park Shell Game: Family's Home Destroyed After Eviction by Fraudulent LLCs"
)

BODY = [
    "Entity bait-and-switch between Cricklewood and Shady Oaks Park MHP LLC.",
    "Sewer violations ignored and eviction executed after rent was paid.",
    "Court clerks blocked filings and labeled counterclaims as moot.",
    "Homes of America links multiple parks to a single owner shell." 
]


def build_press_summary():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading(SUMMARY, 0)
    for para in BODY:
        doc.add_paragraph(para)
    out_path = os.path.join(OUTPUT_DIR, 'press_complaint_summary_shady_oaks.docx')
    doc.save(out_path)
    print(f'Press summary saved to {out_path}')

if __name__ == '__main__':
    build_press_summary()
