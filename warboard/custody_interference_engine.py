import os
import json
from docx import Document
from .svg_builder import generate_svg_warboard
from .svg_motion_binder import bind_motion_links

CUSTODY_EVENTS = [
    {"date": "2024-03-26", "description": "Parenting time blocked"},
    {"date": "2024-04-02", "description": "Filed original custody motion"},
    {"date": "2024-06-01", "description": "FOC refused to enter parenting order"},
    {"date": "2024-08-23", "description": "Case closed despite open motions"},
    {"date": "2025-02-12", "description": "Contempt ruling with no findings"},
]

DOCX_EXPORT = os.path.join('warboard', 'exports', 'CUSTODY_INTERFERENCE_MAP.docx')
SVG_EXPORT = os.path.join('warboard', 'exports', 'CUSTODY_INTERFERENCE_MAP.svg')


def build_custody_warboard():
    os.makedirs(os.path.dirname(DOCX_EXPORT), exist_ok=True)
    doc = Document()
    doc.add_heading('CUSTODY INTERFERENCE MAP', 0)
    for e in CUSTODY_EVENTS:
        doc.add_paragraph(f"{e['date']} - {e['description']}")
    doc.save(DOCX_EXPORT)

    generate_svg_warboard(svg_path=SVG_EXPORT, events=CUSTODY_EVENTS)
    bind_motion_links(svg_path=SVG_EXPORT)
    print(f'Custody interference map generated at {DOCX_EXPORT}')


if __name__ == '__main__':
    build_custody_warboard()
