import os
from docx import Document
from .svg_builder import generate_svg_warboard
from .svg_motion_binder import bind_motion_links

PPO_EVENTS = [
    {"date": "2023-10-15", "description": "Emily Watson filed false welfare call"},
    {"date": "2023-11-10", "description": "Show Cause #1 served (alleged contact)"},
    {"date": "2024-03-26", "description": "Denied parenting time begins"},
    {"date": "2024-04-01", "description": "AppClose log shows mutual contact"},
    {"date": "2024-04-02", "description": "False 'arsenic' photo entered"},
    {"date": "2025-02-12", "description": "Contempt ordered, verbal only (Show Cause #4)"},
    {"date": "2025-03-15", "description": "Job lost after jailing, due to PPO order"},
]

DOCX_EXPORT = os.path.join('warboard', 'exports', 'PPO_WARBOARD.docx')
SVG_EXPORT = os.path.join('warboard', 'exports', 'PPO_WARBOARD.svg')


def build_ppo_warboard():
    """Generate PPO warboard DOCX and SVG from predefined events."""
    os.makedirs(os.path.dirname(DOCX_EXPORT), exist_ok=True)
    doc = Document()
    doc.add_heading('PPO WARBOARD', 0)
    for event in PPO_EVENTS:
        date = event['date']
        desc = event['description']
        doc.add_paragraph(f"{date} - {desc}")
    doc.save(DOCX_EXPORT)

    generate_svg_warboard(events=PPO_EVENTS, svg_path=SVG_EXPORT)
    bind_motion_links(svg_path=SVG_EXPORT)
    print(f'PPO warboard generated at {DOCX_EXPORT}')


if __name__ == '__main__':
    build_ppo_warboard()
