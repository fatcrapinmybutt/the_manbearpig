import json
import os
from docx import Document
from datetime import datetime
"""Utilities for building a Warboard DOCX and SVG and optionally uploading
results to Google Drive. Can be executed as a module or script."""

# Import using an absolute path so execution as a script works as well
from scanner.scan_engine import run_scan
from timeline.builder import build_timeline
from contradictions.contradiction_matrix import detect_contradictions
from warboard.svg_builder import generate_svg_warboard
from warboard.svg_motion_binder import bind_motion_links
from gdrive_sync import upload_to_drive

DOCX_EXPORT = os.path.join('warboard', 'exports', 'SHADY_OAKS_WARBOARD.docx')
SVG_EXPORT = os.path.join('warboard', 'exports', 'SHADY_OAKS_WARBOARD.svg')

TIMELINE_FILE = os.path.join('data', 'timeline.json')
CONTRADICTIONS_FILE = os.path.join('data', 'contradiction_matrix.json')


def build_warboard_docx():
    """Create a DOCX summary of the timeline and contradictions."""
    doc = Document()
    doc.add_heading('SHADY OAKS WARBOARD', 0)

    if os.path.exists(TIMELINE_FILE):
        with open(TIMELINE_FILE, 'r') as f:
            timeline = json.load(f)
        doc.add_heading('Timeline of Events', level=1)
        for event in timeline:
            date = event.get('date', '')[:10]
            desc = event.get('description', '')
            doc.add_paragraph(f"{date} - {desc}")

    if os.path.exists(CONTRADICTIONS_FILE):
        with open(CONTRADICTIONS_FILE, 'r') as f:
            contradictions = json.load(f)
        doc.add_heading('Contradictions', level=1)
        for c in contradictions:
            a = os.path.basename(c.get('file_a', ''))
            b = os.path.basename(c.get('file_b', ''))
            contr = c.get('contradiction', '')
            doc.add_paragraph(f"{a} vs {b}: {contr}")

    os.makedirs(os.path.dirname(DOCX_EXPORT), exist_ok=True)
    doc.save(DOCX_EXPORT)
    print(f"Warboard DOCX saved to {DOCX_EXPORT}")


def deploy_supra_warboard():
    """Run full scan, timeline generation, warboard build, and optional upload."""
    run_scan()
    build_timeline()
    detect_contradictions()
    build_warboard_docx()
    generate_svg_warboard(svg_path=SVG_EXPORT)
    bind_motion_links(svg_path=SVG_EXPORT)
    if os.path.exists('token.json'):
        upload_to_drive(DOCX_EXPORT)
        upload_to_drive(SVG_EXPORT)


if __name__ == '__main__':
    deploy_supra_warboard()
